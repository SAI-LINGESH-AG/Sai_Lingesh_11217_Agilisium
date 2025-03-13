# import streamlit as st
# import google.generativeai as genai
# import json
# import time
# import config

# genai.configure(api_key=config.API_KEY)

# MODEL = "models/gemini-1.5-pro-latest"

# def generate_document(doc_type, details):
#     prompt = f"""
#     Generate a {doc_type} document based on the following details:
#     {details}
#     """
#     model = genai.GenerativeModel(MODEL)
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating document."

# def main():
#     st.title("DocGen")
#     st.markdown("Generate BRD, HLDd, LLDd, and Design Documents")
    
#     doc_type = st.selectbox("Select Document Type:", ["BRD", "HLDd", "LLDd", "Design Document"])
#     details = st.text_area("Enter project details:")
    
#     if st.button("Generate Document"):
#         if details:
#             with st.spinner("Generating document..."):
#                 doc_content = generate_document(doc_type, details)
#                 st.session_state.generated_doc = doc_content
#                 st.success("Document generated successfully!")
#         else:
#             st.warning("Please enter project details.")
    
#     if "generated_doc" in st.session_state:
#         st.text_area("Generated Document:", st.session_state.generated_doc, height=300)
        
#         file_name = f"{doc_type}_Document.txt"
#         st.download_button(
#             label="Download Document",
#             data=st.session_state.generated_doc,
#             file_name=file_name,
#             mime="text/plain"
#         )

# if __name__ == "__main__":
#     main()

################################################################################################################

# import streamlit as st
# import google.generativeai as genai
# from docx import Document
# import io
# import config

# genai.configure(api_key=config.API_KEY)

# MODEL = "models/gemini-1.5-pro-latest"

# def generate_document(doc_type, details):
#     prompt = f"""
#     Generate a {doc_type} document based on the following details:
#     {details}
#     """
#     model = genai.GenerativeModel(MODEL)
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating document."

# def save_as_docx(text, doc_type):
#     doc = Document()
#     doc.add_heading(f"{doc_type} Document", level=1)
#     doc.add_paragraph(text)
    
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# def main():
#     st.title("DocGen")
#     st.markdown("Generate BRD, HLDd, LLDd, and Design Documents")
    
#     doc_type = st.selectbox("Select Document Type:", ["BRD", "HLDd", "LLDd", "Design Document"])
#     details = st.text_area("Enter project details:")
    
#     if st.button("Generate Document"):
#         if details:
#             with st.spinner("Generating document..."):
#                 doc_content = generate_document(doc_type, details)
#                 st.session_state.generated_doc = doc_content
#                 st.success("Document generated successfully!")
#         else:
#             st.warning("Please enter project details.")
    
#     if "generated_doc" in st.session_state:
#         st.text_area("Generated Document:", st.session_state.generated_doc, height=300)
        
#         file_name = f"{doc_type}_Document.docx"
#         docx_file = save_as_docx(st.session_state.generated_doc, doc_type)
        
#         st.download_button(
#             label="Download Document",
#             data=docx_file,
#             file_name=file_name,
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

# if __name__ == "__main__":
#     main()

#################################################################################################

# import streamlit as st
# import google.generativeai as genai
# from docx import Document
# import io
# import config

# # Configure Gemini API
# genai.configure(api_key=config.API_KEY)
# MODEL = "models/gemini-1.5-pro-latest"

# def generate_document(doc_type, details):
#     prompt = f"""
#     Generate a well-structured {doc_type} document with proper formatting, headings, and bullet points based on the following details:
#     {details}
#     """
#     model = genai.GenerativeModel(MODEL)
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating document."

# def save_as_docx(text, doc_type):
#     doc = Document()
#     doc.add_heading(f"{doc_type} Document", level=1)
    
#     # Split text into lines and format properly
#     lines = text.split("\n")
#     for line in lines:
#         if line.strip():
#             if line.startswith("**") and line.endswith("**"):  # Bold headings
#                 doc.add_heading(line.strip("*"), level=2)
#             elif line.startswith("* ") or line.startswith("- "):  # Bullet points
#                 doc.add_paragraph(line.strip("* -"), style="List Bullet")
#             else:
#                 doc.add_paragraph(line)
    
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# def main():
#     st.set_page_config(page_title="Document Generator", layout="wide")
#     st.title("📄 DocGen")
#     st.markdown("### Generate BRD, HLDd, LLDd, and Design Documents with AI-powered assistance.")
    
#     doc_type = st.selectbox("📌 Select Document Type:", ["BRD", "HLDd", "LLDd", "Design Document"])
#     details = st.text_area("✍️ Enter project details:", height=150)
    
#     if st.button("🚀 Generate Document"):
#         if details:
#             with st.spinner("Generating document..."):
#                 doc_content = generate_document(doc_type, details)
#                 st.session_state.generated_doc = doc_content
#                 st.success("✅ Document generated successfully!")
#         else:
#             st.warning("⚠️ Please enter project details.")
    
#     if "generated_doc" in st.session_state:
#         st.subheader("📑 Generated Document Preview:")
#         st.text_area("", st.session_state.generated_doc, height=300)
        
#         file_name = f"{doc_type}_Document.docx"
#         docx_file = save_as_docx(st.session_state.generated_doc, doc_type)
        
#         st.download_button(
#             label="📥 Download Document",
#             data=docx_file,
#             file_name=file_name,
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

# if __name__ == "__main__":
#     main()

#################################################################################################################

# import streamlit as st
# import google.generativeai as genai
# from docx import Document
# import io
# import config

# # Configure Gemini API
# genai.configure(api_key=config.API_KEY)
# MODEL = "models/gemini-1.5-pro-latest"

# def generate_document(doc_type, details):
#     prompt = f"""
#     Generate a well-structured {doc_type} document with proper formatting, headings, numbered sections, and bullet points based on the following details:
#     {details}
#     """
#     model = genai.GenerativeModel(MODEL)
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating document."

# def save_as_docx(text, doc_type):
#     doc = Document()
#     doc.add_heading(f"{doc_type} Document", level=1)
    
#     lines = text.split("\n")
#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue
        
#         if line[0].isdigit() and "." in line:  # Numbered sections
#             doc.add_heading(line, level=2)
#         elif line.startswith("**") and line.endswith("**"):  # Bold headings
#             doc.add_heading(line.strip("*"), level=3)
#         elif line.startswith("- ") or line.startswith("* "):  # Bullet points
#             doc.add_paragraph(line.strip("- *"), style="List Bullet")
#         else:
#             doc.add_paragraph(line)
    
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# def main():
#     st.set_page_config(page_title="Document Generator", layout="wide")
#     st.title("📄 DocGen")
#     st.markdown("### Generate BRD, HLDd, LLDd, and Design Documents with AI-powered assistance.")
    
#     doc_type = st.selectbox("📌 Select Document Type:", ["BRD", "HLDd", "LLDd", "Design Document"])
#     details = st.text_area("✍️ Enter project details:", height=150)
    
#     if st.button("🚀 Generate Document"):
#         if details:
#             with st.spinner("Generating document..."):
#                 doc_content = generate_document(doc_type, details)
#                 st.session_state.generated_doc = doc_content
#                 st.success("✅ Document generated successfully!")
#         else:
#             st.warning("⚠️ Please enter project details.")
    
#     if "generated_doc" in st.session_state:
#         st.subheader("📑 Generated Document Preview:")
#         st.text_area("", st.session_state.generated_doc, height=300)
        
#         file_name = f"{doc_type}_Document.docx"
#         docx_file = save_as_docx(st.session_state.generated_doc, doc_type)
        
#         st.download_button(
#             label="📥 Download Document",
#             data=docx_file,
#             file_name=file_name,
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

# if __name__ == "__main__":
#     main()

####################################################################################################

import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import config

# Configure Gemini API
genai.configure(api_key=config.API_KEY)
MODEL = "models/gemini-1.5-pro-latest"

def generate_document(doc_type, details):
    prompt = f"""
    Generate a well-structured {doc_type} document with proper formatting, headings, numbered sections, and bullet points based on the following details:
    {details}
    Ensure that no metadata such as 'Version', 'Date', or 'Prepared by' is included.
    """
    model = genai.GenerativeModel(MODEL)
    response = model.generate_content(prompt)
    return response.text if response else "Error generating document."

def save_as_docx(text, doc_type):
    doc = Document()
    doc.add_heading(f"{doc_type} Document", level=1)
    
    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line[0].isdigit() and "." in line:  # Numbered sections
            doc.add_heading(line, level=2)
        elif line.startswith("**") and line.endswith("**"):  # Bold headings
            doc.add_heading(line.strip("*"), level=3)
        elif line.startswith("- ") or line.startswith("* "):  # Bullet points
            doc.add_paragraph(line.strip("- *"), style="List Bullet")
        else:
            doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(page_title="Document Generator", layout="wide")
    st.title("📄 DocGen")
    st.markdown("### Generate BRD, HLDd, LLDd, and Design Documents with AI-powered assistance.")
    
    doc_type = st.selectbox("📌 Select Document Type:", ["BRD", "HLDd", "LLDd", "Design Document"])
    details = st.text_area("✍️ Enter project details:", height=150)
    
    if st.button("🚀 Generate Document"):
        if details:
            with st.spinner("Generating document..."):
                doc_content = generate_document(doc_type, details)
                st.session_state.generated_doc = doc_content
                st.success("✅ Document generated successfully!")
        else:
            st.warning("⚠️ Please enter project details.")
    
    if "generated_doc" in st.session_state:
        st.subheader("📑 Generated Document Preview:")
        st.text_area("", st.session_state.generated_doc, height=300)
        
        file_name = f"{doc_type}_Document.docx"
        docx_file = save_as_docx(st.session_state.generated_doc, doc_type)
        
        st.download_button(
            label="📥 Download Document",
            data=docx_file,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()

###########################################################################################################

# import google.generativeai as genai
# import config

# # Configure Gemini API
# genai.configure(api_key=config.API_KEY)
# MODEL = "models/gemini-1.5-pro-latest"

################################################################################################

# import streamlit as st
# import google.generativeai as genai
# from docx import Document
# import io
# import config

# # Configure Gemini API
# genai.configure(api_key=config.API_KEY)
# MODEL = "models/gemini-1.5-pro-latest"

# def generate_document(doc_type, details):
#     prompt = f"""
#     Generate a {doc_type} document in plain text format. Do not include any formatting such as bold text, numbered sections, or bullet points.
#     Just provide the content in a simple text format.

#     Details:
#     {details}
#     """
#     model = genai.GenerativeModel(MODEL)
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating document."

# def save_as_docx(text):
#     doc = Document()
#     doc.add_paragraph(text)  
    
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# def main():
#     st.set_page_config(page_title="Document Generator", layout="wide")
#     st.title("📄 DocGen")

#     doc_type = st.selectbox("📌 Select Document Type:", ["BRD", "HLD", "LLD", "Design Document"])
#     details = st.text_area("✍️ Enter project details:", height=150)

#     if st.button("🚀 Generate Document"):
#         if details:
#             with st.spinner("Generating document..."):
#                 doc_content = generate_document(doc_type, details)
#                 st.session_state.generated_doc = doc_content
#                 st.success("✅ Document generated successfully!")
#         else:
#             st.warning("⚠️ Please enter project details.")

#     if "generated_doc" in st.session_state:
#         st.subheader("📑 Generated Document Preview:")
#         st.text_area("", st.session_state.generated_doc, height=300)

#         file_name = f"{doc_type}_Document.docx"
#         docx_file = save_as_docx(st.session_state.generated_doc)

#         st.download_button(
#             label="📥 Download Word Document",
#             data=docx_file,
#             file_name=file_name,
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

# if __name__ == "__main__":
#     main()

################################################################################

# import config
# import streamlit as st
# import json
# import docx
# import pdfkit
# import yaml
# import google.generativeai as genai

# genai.configure(api_key=config.API_KEY)

# # Function to generate document content using Gemini API
# def generate_content(prompt):
#     model = genai.GenerativeModel("gemini-1.5-flash-latest")  # Fast response model
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating content."

# # Function to save as DOCX
# def save_as_docx(content, filename):
#     doc = docx.Document()
#     doc.add_paragraph(content)
#     doc.save(filename)
#     return filename

# # Function to save as PDF
# def save_as_pdf(content, filename):
#     pdfkit.from_string(content, filename)
#     return filename

# # Function to save as JSON
# def save_as_json(content, filename):
#     with open(filename, "w") as f:
#         json.dump({"content": content}, f)
#     return filename

# # Function to save as YAML
# def save_as_yaml(content, filename):
#     with open(filename, "w") as f:
#         yaml.dump({"content": content}, f)
#     return filename

# # Streamlit UI
# st.title("AI-Powered Document Generator")

# # User input
# prompt = st.text_area("Enter document description:")
# file_format = st.selectbox("Choose file format:", ["DOCX", "PDF", "JSON", "YAML"])

# if st.button("Generate Document"):
#     if prompt:
#         generated_text = generate_content(prompt)
#         st.text_area("Generated Content:", generated_text, height=300)
        
#         # Save file in selected format
#         filename = f"generated_document.{file_format.lower()}"
#         if file_format == "DOCX":
#             save_as_docx(generated_text, filename)
#         elif file_format == "PDF":
#             save_as_pdf(generated_text, filename)
#         elif file_format == "JSON":
#             save_as_json(generated_text, filename)
#         elif file_format == "YAML":
#             save_as_yaml(generated_text, filename)
        
#         st.success(f"Document generated successfully!")
#         st.download_button(label="Download Document", data=open(filename, "rb").read(), file_name=filename)
#     else:
#         st.warning("Please enter a description to generate the document.")

#########################################################################################

# import streamlit as st
# import google.generativeai as genai
# from docx import Document
# from io import BytesIO
# import config

# # Configure Gemini API
# genai.configure(api_key=config.API_KEY)

# # Function to generate BRD using Gemini API
# def generate_brd(prompt, model="models/gemini-1.5-pro-latest"):
#     model = genai.GenerativeModel(model)
#     response = model.generate_content(prompt)
#     return response.text if response else "Error generating BRD."

# # Function to create a Word document
# def create_word_doc(content):
#     doc = Document()
#     doc.add_heading("Business Requirements Document (BRD)", level=1)
#     doc.add_paragraph(content)
    
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# # Streamlit UI
# st.title("BRD DocGen")
# st.write("Enter your project details in plain English, and this app will generate a Business Requirements Document (BRD).")

# # User Input
# user_input = st.text_area("Enter your project description:")
# selected_model = st.selectbox("Select Gemini Model:", [
#     "models/gemini-1.5-pro-latest",
#     "models/gemini-1.5-flash-latest",
#     "models/gemini-2.0-flash",
#     "models/gemini-2.0-pro-exp"
# ])

# generate_button = st.button("Generate BRD")

# if generate_button and user_input:
#     with st.spinner("Generating BRD..."):
#         brd_text = generate_brd(user_input, selected_model)
#         word_file = create_word_doc(brd_text)
        
#         st.subheader("Generated BRD:")
#         st.text_area("BRD Content", brd_text, height=300)
#         st.download_button(
#             label="Download BRD as Word File",
#             data=word_file,
#             file_name="Business_Requirements_Document.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )